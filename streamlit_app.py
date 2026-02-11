
import io
import math
import base64, mimetypes
import streamlit.components.v1 as components
from datetime import date
from pathlib import Path
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import streamlit as st
import os
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from docx import Document

st.set_page_config(page_title="WorkSmart Hub – Executive Demo", page_icon="🚜", layout="wide")

# ==============================
# CSS & Background Helpers
# ==============================

def inject_css(path: str = "style.css"):
    """Read a CSS file and inject it into the app."""
    p = Path(path)
    if p.exists():
        css = p.read_text(encoding="utf-8")
        st.markdown(f"<style>{css}</style>", unsafe_allow_html=True)
    else:
        st.warning(f"Missing CSS theme at {path}. The app will still run.")


def apply_tab_background(image_path: str):
    """Embed a background image as data-URI for reliability on Streamlit (if file exists)."""
    p = Path(image_path)
    if not p.exists():
        st.warning(f"Background not found: {image_path}")
        return
    mime = mimetypes.guess_type(p.name)[0] or "image/jpeg"
    data64 = base64.b64encode(p.read_bytes()).decode("ascii")
    st.markdown(
        f"""
        <style>
        .stApp {{
            background: url(data:{mime};base64,{data64}) no-repeat center center fixed !important;
            background-size: cover !important;
        }}
        </style>
        """,
        unsafe_allow_html=True,
    )


def render_stage_with_overlay(bg_path: str, blocks: dict, values: dict, aspect_ratio=(3, 2)):
    """
    Render a fixed-aspect stage with bg image + overlay blocks positioned by percentages.
    blocks: { key: {"x_pct": float, "y_pct": float, "w_pct": float, "align": "center|left|right", "font_px": int} }
    values: { key: "text" }
    aspect_ratio: (W, H) of the artwork; current assess image is ~3:2.
    """
    p = Path(bg_path)
    if not p.exists():
        st.warning(f"Background not found: {bg_path}")
        return

    w, h = aspect_ratio
    pad_top_pct = (h / w) * 100.0

    mime = mimetypes.guess_type(p.name)[0] or "image/jpeg"
    data64 = base64.b64encode(p.read_bytes()).decode("ascii")

    # Build overlay HTML blocks
    parts = []
    for key, cfg in blocks.items():
        x = float(cfg.get("x_pct", 50))
        y = float(cfg.get("y_pct", 50))
        w_pct = float(cfg.get("w_pct", 12))
        align = cfg.get("align", "center")
        font_px = int(cfg.get("font_px", 22))
        val = values.get(key, "")
        parts.append(
            f'<div class="ws-overlay-block" style="left:{x}%; top:{y}%; width:{w_pct}%; text-align:{align};">'
            f'<div class="value" style="font-size:{font_px}px;">{val}</div>'
            f'</div>'
        )
    overlay_html = "".join(parts)

    html = f"""
    <div class="ws-stage">
      <div class="ws-stage-inner" style="padding-top:{pad_top_pct}%;
           background-image:url(data:{mime};base64,{data64});">
        <div class="ws-overlay">{overlay_html}</div>
      </div>
    </div>
    """
    components.html(html, height=0)


# ==============================
# Overlay Engine (align live values to art slots)
# ==============================

ASSESS_LAYOUT = {
    # Top row (left → right)
    "fleet":      {"x_pct": 22.0, "y_pct": 14.0, "w_pct": 12.0, "align": "center", "font_px": 20},
    "acres":      {"x_pct": 42.0, "y_pct": 14.0, "w_pct": 12.0, "align": "center", "font_px": 20},
    "util_rate":  {"x_pct": 61.5, "y_pct": 14.0, "w_pct": 14.0, "align": "center", "font_px": 20},
    "idle_time":  {"x_pct": 80.0, "y_pct": 14.0, "w_pct": 12.0, "align": "center", "font_px": 20},

    # Bottom tiles (left → right)
    "maint_cost": {"x_pct": 14.0, "y_pct": 86.0, "w_pct": 16.0, "align": "center", "font_px": 24},
    "downtime":   {"x_pct": 32.0, "y_pct": 86.0, "w_pct": 16.0, "align": "center", "font_px": 24},
    "safety":     {"x_pct": 50.0, "y_pct": 86.0, "w_pct": 16.0, "align": "center", "font_px": 24},
    "rework":     {"x_pct": 68.0, "y_pct": 86.0, "w_pct": 16.0, "align": "center", "font_px": 24},
    "fuel":       {"x_pct": 86.0, "y_pct": 86.0, "w_pct": 12.0, "align": "center", "font_px": 24},
}

ASSESS_OVERLAY = {
    "fleet": {"x_vw": 22.0, "y_vh": 10.2, "w_vw": 12.0, "align": "center", "cls": "hud-top"},
    "acres": {"x_vw": 42.0, "y_vh": 10.2, "w_vw": 12.0, "align": "center", "cls": "hud-top"},
    "util_rate": {"x_vw": 61.0, "y_vh": 10.2, "w_vw": 14.0, "align": "center", "cls": "hud-top"},
    "idle_time": {"x_vw": 80.0, "y_vh": 10.2, "w_vw": 12.0, "align": "center", "cls": "hud-top"},
    "maint_cost": {"x_vw": 14.0, "y_vh": 84.0, "w_vw": 16.0, "align": "center", "cls": "hud-bottom"},
    "downtime": {"x_vw": 32.0, "y_vh": 84.0, "w_vw": 16.0, "align": "center", "cls": "hud-bottom"},
    "safety": {"x_vw": 50.0, "y_vh": 84.0, "w_vw": 16.0, "align": "center", "cls": "hud-bottom"},
    "rework": {"x_vw": 68.0, "y_vh": 84.0, "w_vw": 16.0, "align": "center", "cls": "hud-bottom"},
    "fuel": {"x_vw": 86.0, "y_vh": 84.0, "w_vw": 10.0, "align": "center", "cls": "hud-bottom"},
}

if "overlay_offsets" not in st.session_state:
    st.session_state.overlay_offsets = {}


def render_overlay(blocks: dict, values: dict, show_guides: bool = False):
    """Render absolutely-positioned live values over the background image."""
    # Build compact HTML (no leading spaces/newlines) so nothing gets escaped.
    parts = ['<div class="ws-overlay">']
    for key, cfg in blocks.items():
        x = cfg.get("x_vw", 0)
        y = cfg.get("y_vh", 0)
        w = cfg.get("w_vw", 10)
        align = cfg.get("align", "left")
        cls = cfg.get("cls", "")

        # Per-slot nudge offsets (in pixels)
        dx = st.session_state.overlay_offsets.get(f"{key}_dx", 0)
        dy = st.session_state.overlay_offsets.get(f"{key}_dy", 0)

        # Optional alignment guides (dashed border + subtle background)
        guide_css = ("border:1px dashed rgba(255,255,255,.15); background: rgba(0,0,0,.08);"
                     if show_guides else "")

        # Single-line, compact HTML avoids Markdown "code block" escaping
        parts.append(
            f'<div class="ws-overlay-block {cls}" '
            f'style="left:{x}vw; top:{y}vh; width:{w}vw; '
            f'transform:translate({dx}px,{dy}px); text-align:{align}; {guide_css}">'
            f'<div class="value">{values.get(key, "")}</div>'
            f'</div>'
        )
    parts.append("</div>")
    html = "".join(parts)

    # Render as raw HTML (no Markdown) so tags never get escaped
    components.html(html, height=0)

# ==============================
# ROI / P&L utilities
# ==============================
def currency(x):
    try:
        return f"${x:,.0f}"
    except Exception:
        return str(x)

def pct(x):
    try:
        return f"{x*100:.0f}%"
    except Exception:
        return str(x)

def init_defaults():
    defaults = {
        "acres": 1500,
        "crops": "Row crops",
        "machines": {"Tractors": 6, "RTVs": 4, "Construction": 2},
        "util_hours_per_machine": 900,
        "labor_rate": 28.0,
        "fuel_price": 3.75,
        "baseline_idle_pct": 0.22,
        "baseline_downtime_hours_month": 14.0,
        "baseline_maint_cost_year": 48000.0,
        "safety_incidents_year": 2.0,
        "rework_pct": 0.06,
        "yield_uplift_potential": 0.02,
        "telemetry_attach_rate": 0.75,
        "worksmart_price_per_machine_month": 35.0,
        "hardware_kit_price": 450.0,
        "hardware_gross_margin": 0.35,
        "idle_reduction_pct": 0.35,
        "downtime_reduction_pct": 0.25,
        "maint_cost_reduction_pct": 0.12,
        "safety_reduction_pct": 0.30,
        "rework_reduction_pct": 0.30,
        "labor_productivity_gain_pct": 0.05,
        "autosteer_overlap_reduction_pct": 0.12,
        "autosteer_yield_consistency_pct": 0.005,
        "autonomy_productivity_gain_pct": 0.25,
        "autonomy_supervision_gain_pct": 0.20,
        "autonomy_safety_gain_pct": 0.10,
        "discount_rate": 0.12,
        "time_horizon_years": 3,
        "base_deployed_fleet": 10000,
        "attach_rate": 0.35,
        "attach_growth_yoy": 0.20,
        "churn_rate": 0.06,
        "cac_per_attach": 120.0,
    }
    if "inputs" not in st.session_state:
        st.session_state.inputs = defaults.copy()
    else:
        # Ensure any new defaults are present without overwriting current user values
        for k, v in defaults.items():
            if k not in st.session_state.inputs:
                st.session_state.inputs[k] = v

def total_machines(inputs):
    return int(sum(inputs["machines"].values()))

def estimate_fuel_use(inputs):
    gph = {"Tractors": 3.8, "RTVs": 1.4, "Construction": 4.5}
    hours = inputs["util_hours_per_machine"]
    total_gal = 0.0
    for k, v in inputs["machines"].items():
        total_gal += gph.get(k, 3.0) * hours * v
    return total_gal

def compute_roi(inputs, actions):
    n = total_machines(inputs)
    attach = min(max(inputs["telemetry_attach_rate"], 0.0), 1.0)
    attached_machines = int(round(n * attach))

    total_hours = n * inputs["util_hours_per_machine"]

    baseline_fuel_gal = estimate_fuel_use(inputs)
    fuel_price = inputs["fuel_price"]
    labor_rate = inputs["labor_rate"]

    idle_reduction = inputs["idle_reduction_pct"] if actions.get("reduce_idling") else 0.0
    downtime_reduction = inputs["downtime_reduction_pct"] if actions.get("predictive_maint") else 0.0
    maint_cost_reduction = inputs["maint_cost_reduction_pct"] if actions.get("predictive_maint") else 0.0
    safety_reduction = inputs["safety_reduction_pct"] if actions.get("safety_monitoring") else 0.0
    rework_reduction = inputs["rework_reduction_pct"] if actions.get("guided_ops") else 0.0
    labor_productivity = inputs["labor_productivity_gain_pct"] if actions.get("guided_ops") else 0.0
    yield_uplift = inputs["yield_uplift_potential"] if actions.get("optimize_ops") else 0.0

    autosteer = actions.get("autosteer", False)
    autonomy = actions.get("autonomy", False)

    idle_share = inputs["baseline_idle_pct"]
    fuel_idle = baseline_fuel_gal * idle_share
    fuel_saved_gal = fuel_idle * idle_reduction
    fuel_savings = fuel_saved_gal * fuel_price

    downtime_hours_year = inputs["baseline_downtime_hours_month"] * 12
    downtime_hours_saved = downtime_hours_year * downtime_reduction
    downtime_value = downtime_hours_saved * labor_rate

    maint_savings = inputs["baseline_maint_cost_year"] * maint_cost_reduction

    avg_cost_per_incident = 15000
    safety_saved = inputs["safety_incidents_year"] * safety_reduction * avg_cost_per_incident

    rework_hours = total_hours * inputs["rework_pct"]
    rework_saved = rework_hours * rework_reduction * labor_rate

    productivity_value = total_hours * labor_rate * labor_productivity

    revenue_per_acre = 350
    yield_uplift_value = inputs["acres"] * revenue_per_acre * yield_uplift

    autosteer_fuel = 0.0
    autosteer_yield = 0.0
    if autosteer:
        non_idle_fuel_gal = baseline_fuel_gal * (1 - idle_share)
        autosteer_fuel = non_idle_fuel_gal * inputs["autosteer_overlap_reduction_pct"] * fuel_price
        autosteer_yield = inputs["acres"] * revenue_per_acre * inputs["autosteer_yield_consistency_pct"]

    autonomy_prod = 0.0
    autonomy_supervision = 0.0
    autonomy_safety = 0.0
    if autonomy:
        autonomy_prod = total_hours * labor_rate * inputs["autonomy_productivity_gain_pct"]
        autonomy_supervision = total_hours * labor_rate * inputs["autonomy_supervision_gain_pct"]
        autonomy_safety = inputs["safety_incidents_year"] * avg_cost_per_incident * inputs["autonomy_safety_gain_pct"]

    gross_annual_benefit = (
        fuel_savings
        + downtime_value
        + maint_savings
        + safety_saved
        + rework_saved
        + productivity_value
        + yield_uplift_value
        + autosteer_fuel
        + autosteer_yield
        + autonomy_prod
        + autonomy_supervision
        + autonomy_safety
    )

    subscription_cost = attached_machines * inputs["worksmart_price_per_machine_month"] * 12
    hardware_cost = attached_machines * inputs["hardware_kit_price"]

    net_annual_benefit = gross_annual_benefit - subscription_cost

    r = inputs["discount_rate"]
    years = inputs["time_horizon_years"]
    cash_flows = [-hardware_cost] + [net_annual_benefit for _ in range(years)]
    npv = sum(cf / ((1 + r) ** t) for t, cf in enumerate(cash_flows))

    monthly_net = net_annual_benefit / 12 if net_annual_benefit != 0 else 0.00001
    payback_months = hardware_cost / monthly_net if monthly_net > 0 else float("inf")

    hardware_margin = inputs["hardware_gross_margin"]
    kubota_hw_gm = hardware_cost * hardware_margin
    kubota_sw_revenue = subscription_cost
    kubota_total_revenue_yr1 = kubota_sw_revenue + hardware_cost

    warranty_claim_reduction = 0.08 if actions.get("predictive_maint") else 0.0
    warranty_savings = 1200 * attached_machines * warranty_claim_reduction * 0.5

    parts_service_uplift = 180 * attached_machines

    components = {
        "Fuel savings (idle)": fuel_savings,
        "Downtime reduction": downtime_value,
        "Maintenance savings": maint_savings,
        "Safety savings": safety_saved,
        "Rework savings": rework_saved,
        "Productivity gain": productivity_value,
        "Yield uplift": yield_uplift_value,
        "AutoSteer: fuel/overlap": autosteer_fuel,
        "AutoSteer: yield consistency": autosteer_yield,
        "Autonomy: productivity": autonomy_prod,
        "Autonomy: supervision": autonomy_supervision,
        "Autonomy: safety": autonomy_safety,
        "Subscription cost": -subscription_cost,
    }

    return {
        "attached_machines": attached_machines,
        "gross_annual_benefit": gross_annual_benefit,
        "subscription_cost": subscription_cost,
        "hardware_cost": hardware_cost,
        "net_annual_benefit": net_annual_benefit,
        "npv": npv,
        "payback_months": payback_months,
        "kubota_total_revenue_yr1": kubota_total_revenue_yr1,
        "kubota_hw_gm": kubota_hw_gm,
        "kubota_sw_revenue": kubota_sw_revenue,
        "warranty_savings": warranty_savings,
        "parts_service_uplift": parts_service_uplift,
        "components": components,
    }

def enterprise_pnl(inputs):
    price = inputs["worksmart_price_per_machine_month"]
    hw_price = inputs["hardware_kit_price"]
    hw_margin = inputs["hardware_gross_margin"]
    attach = inputs["attach_rate"]
    base_fleet = inputs["base_deployed_fleet"]
    growth = inputs["attach_growth_yoy"]
    churn = inputs["churn_rate"]
    cac = inputs["cac_per_attach"]

    years = [1, 2, 3]
    data = []

    attached_prev = base_fleet * attach
    for y in years:
        gross_adds = attached_prev * growth
        churned = attached_prev * churn
        net_adds = max(gross_adds - churned, 0)
        attached = attached_prev + net_adds

        arr = attached * price * 12
        hw_rev = net_adds * hw_price
        hw_gm = hw_rev * hw_margin
        cac_cost = net_adds * cac

        total_rev = arr + hw_rev
        gm_est = hw_gm + arr * 0.8  # assume 80% SW GM proxy
        ebitda_proxy = gm_est - cac_cost

        data.append(
            {
                "Year": f"Y{y}",
                "Attached (eop)": int(attached),
                "ARR": arr,
                "Hardware Rev": hw_rev,
                "Hardware GM": hw_gm,
                "CAC": cac_cost,
                "Total Rev": total_rev,
                "GM (proxy)": gm_est,
                "EBITDA (proxy)": ebitda_proxy,
            }
        )

        attached_prev = attached

    return pd.DataFrame(data)

# ==============================
# Init & Tabs
# ==============================
init_defaults()

if "actions" not in st.session_state:
    st.session_state.actions = {
        "reduce_idling": True,
        "predictive_maint": True,
        "safety_monitoring": True,
        "guided_ops": True,
        "optimize_ops": False,
        "autosteer": True,
        "autonomy": False,
    }

inject_css("style.css")

oview, assess, analyze, act, roi, kubota, roadmap, hardware, anim, dealer, console, export = st.tabs(
    [
        "Overview",
        "Assess",
        "Analyze",
        "Act",
        "ROI & Business Case",
        "Kubota Impact",
        "Roadmap",
        "Hardware Vision",
        "Autonomy Animation",
        "Dealer View",
        "Supervision Console",
        "Export",
    ]
)

with oview:
    apply_tab_background("assets/backgrounds/assess_bg.jpg")
    st.markdown("""
    ### Overview
    """)
    st.caption("A retro–future Kubota HUD showcasing Assess → Analyze → Act, ROI, and enterprise scale.")

    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Connected fleet (demo)", 24)
    with col2:
        st.metric("Avg idle share", pct(st.session_state.inputs["baseline_idle_pct"]))
    with col3:
        st.metric("Downtime (hrs/mo)", f"{st.session_state.inputs['baseline_downtime_hours_month']:.0f}")


with assess:
    # (No need to call apply_tab_background here anymore)
    st.markdown("### Assess – Baseline the Operation")

    # --- (Optional) Alignment tools you already have can remain below ---

    # Build the live values for the overlay from current inputs
    i = st.session_state.inputs
    util_rate = min(i["util_hours_per_machine"] / 2000.0, 1.0)
    values = {
        "fleet": f"{total_machines(i)}",
        "acres": f"{i['acres']:,}",
        "util_rate": f"{util_rate*100:.0f}%",
        "idle_time": f"{i['baseline_idle_pct']*100:.0f}%",
        "maint_cost": currency(i["baseline_maint_cost_year"]),
        "downtime": f"{i['baseline_downtime_hours_month']*12:.0f}",
        "safety": f"{i['safety_incidents_year']:.0f}",
        "rework": f"{i['rework_pct']*100:.0f}%",
        "fuel": f"${i['fuel_price']:.2f}",
    }

    # >>> THIS is the one call that renders the clean background + overlay <<<
    render_stage_with_overlay(
        bg_path="assets/backgrounds/assess_bg.jpg",  # you kept this name; perfect
        blocks=ASSESS_LAYOUT,
        values=values,
        aspect_ratio=(3, 2),  # current Assess artwork is ~3:2; update if you export differently
    )

    # --- Your interactive inputs and sliders can remain below the stage ---
    st.markdown("---")
    c1, c2, c3 = st.columns(3)
    with c1:
        i["acres"] = st.number_input("Acres managed", 1, 200000, int(i["acres"]))
        i["crops"] = st.text_input("Crop type(s)", str(i["crops"]))
        i["util_hours_per_machine"] = st.number_input(
            "Utilization (hrs/machine/yr)", 0, 3000, int(i["util_hours_per_machine"])
        )
    with c2:
        st.markdown("**Fleet size**")
        for k in list(i["machines"].keys()):
            i["machines"][k] = st.number_input(k, 0, 200, int(i["machines"][k]))
        i["telemetry_attach_rate"] = st.slider("Attach rate", 0.0, 1.0, float(i["telemetry_attach_rate"]))
    with c3:
        i["labor_rate"] = st.number_input("Labor rate ($/hr)", 0.0, 500.0, float(i["labor_rate"]))
        i["fuel_price"] = st.number_input("Fuel price ($/gal)", 0.0, 20.0, float(i["fuel_price"]))
        i["baseline_idle_pct"] = st.slider("Idle share", 0.0, 0.7, float(i["baseline_idle_pct"]))

    st.markdown("---")
    c4, c5, c6 = st.columns(3)
    with c4:
        i["baseline_downtime_hours_month"] = st.number_input(
            "Unplanned downtime (hrs/mo)", 0.0, 1000.0, float(i["baseline_downtime_hours_month"])
        )
        i["baseline_maint_cost_year"] = st.number_input(
            "Maintenance cost ($/yr)", 0.0, 1e7, float(i["baseline_maint_cost_year"])
        )
    with c5:
        i["safety_incidents_year"] = st.number_input(
            "Safety incidents (per yr)", 0.0, 100.0, float(i["safety_incidents_year"])
        )
        i["rework_pct"] = st.slider("Rework (% of labor)", 0.0, 0.5, float(i["rework_pct"]))
    with c6:
        i["yield_uplift_potential"] = st.slider("Potential yield uplift", 0.0, 0.2, float(i["yield_uplift_potential"]))

with analyze:
    apply_tab_background("assets/backgrounds/analyze_bg.jpg")
    st.markdown("""
    ### Analyze – Insights
    """)
    i = st.session_state.inputs
    insights = []
    if i["baseline_idle_pct"] >= 0.2:
        insights.append("High idling detected; operator coaching & auto–idle rules")
    if i["baseline_downtime_hours_month"] >= 12:
        insights.append("Elevated unplanned downtime; predictive maintenance recommended")
    if i["rework_pct"] >= 0.05:
        insights.append("Material rework; guided workflows & checklists")
    if i["safety_incidents_year"] >= 2:
        insights.append("Safety incident rate high; monitoring & alerts advised")
    if i["yield_uplift_potential"] >= 0.02:
        insights.append("Yield optimization potential; variable–rate & route planning")

    if insights:
        for msg in insights:
            st.markdown(f"- {msg}")
    else:
        st.success("No critical issues detected. Optimization potential remains for fuel and productivity.")

with act:
    apply_tab_background("assets/backgrounds/act_bg.jpg")
    st.markdown("""
    ### Act – Next Best Actions
    """)
    c1, c2 = st.columns(2)
    with c1:
        st.session_state.actions["reduce_idling"] = st.checkbox("Operator coaching + auto–idle rules", True)
        st.session_state.actions["predictive_maint"] = st.checkbox("Predictive maintenance & parts pre–positioning", True)
        st.session_state.actions["safety_monitoring"] = st.checkbox("Safety monitoring & alerts", True)
        st.session_state.actions["autosteer"] = st.checkbox("AutoSteer (assist)", True)
    with c2:
        st.session_state.actions["guided_ops"] = st.checkbox("Guided workflows & digital checklists", True)
        st.session_state.actions["optimize_ops"] = st.checkbox("Optimization (path, variable–rate, sequencing)", False)
        st.session_state.actions["autonomy"] = st.checkbox("Autonomy (supervised)", False)

with roi:
    apply_tab_background("assets/backgrounds/roi_bg.jpg")
    st.markdown("""
    ### Customer ROI & Business Case
    """)
    actions = st.session_state.actions
    results = compute_roi(st.session_state.inputs, actions)

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Machines attached", results["attached_machines"])
    c2.metric("Gross annual benefit", currency(results["gross_annual_benefit"]))
    c3.metric("Subscription cost", currency(results["subscription_cost"]))
    c4.metric("Net annual benefit", currency(results["net_annual_benefit"]))

    c5, c6, c7 = st.columns(3)
    payback_txt = f"{results['payback_months']:.1f} months" if math.isfinite(results['payback_months']) else "> horizon"
    c5.metric("Payback", payback_txt)
    c6.metric("NPV", currency(results["npv"]))
    roi_pct = (results["net_annual_benefit"] - 0.00001) / max(results["subscription_cost"], 1)
    c7.metric("ROI vs. subscription", pct(roi_pct))

    st.markdown(" ")
    comps = results["components"]
    labels = list(comps.keys())
    values = [comps[k] for k in labels]

    fig, ax = plt.subplots(figsize=(9, 4))
    colors = ["#2ca02c" if v >= 0 else "#d62728" for v in values]
    bars = ax.barh(labels, values, color=colors)
    ax.axvline(0, color="black", linewidth=0.8)
    ax.set_xlabel("$ per year")
    for bar, v in zip(bars, values):
        ax.text(
            bar.get_width() + (2000 if v >= 0 else -2000),
            bar.get_y() + bar.get_height() / 2,
            currency(v),
            va="center",
            ha="left" if v >= 0 else "right",
            fontsize=9,
        )
    st.pyplot(fig)

with kubota:
    apply_tab_background("assets/backgrounds/kubota_bg.jpg")
    st.markdown("""
    ### Kubota Impact – 3–Year P&L
    """)
    inp = st.session_state.inputs
    c1, c2, c3 = st.columns(3)
    with c1:
        inp["base_deployed_fleet"] = st.number_input("Base deployed fleet", 1000, 1_000_000, int(inp["base_deployed_fleet"]))
        inp["attach_rate"] = st.slider("Attach rate (current)", 0.05, 1.0, float(inp["attach_rate"]))
    with c2:
        inp["attach_growth_yoy"] = st.slider("Attach growth YoY", 0.0, 1.0, float(inp["attach_growth_yoy"]))
        inp["churn_rate"] = st.slider("Churn rate", 0.0, 0.5, float(inp["churn_rate"]))
    with c3:
        inp["cac_per_attach"] = st.number_input("CAC per new attach ($)", 0.0, 5000.0, float(inp["cac_per_attach"]))

    df = enterprise_pnl(inp)
    st.dataframe(
        df.style.format(
            {
                "ARR": "${:,.0f}",
                "Hardware Rev": "${:,.0f}",
                "Hardware GM": "${:,.0f}",
                "CAC": "${:,.0f}",
                "Total Rev": "${:,.0f}",
                "GM (proxy)": "${:,.0f}",
                "EBITDA (proxy)": "${:,.0f}",
            }
        )
    )

    fig, ax = plt.subplots(figsize=(8, 3))
    ax.plot(df["Year"], df["ARR"], marker="o", label="ARR")
    ax.plot(df["Year"], df["Hardware Rev"], marker="o", label="Hardware Rev")
    ax.set_ylabel("$")
    ax.legend()
    st.pyplot(fig)

with roadmap:
    apply_tab_background("assets/backgrounds/roadmap_bg.jpg")
    st.markdown("""
    ### Roadmap – Assist → Autonomy
    """)
    col1, col2 = st.columns(2)
    with col1:
        st.markdown(
            """
            **Phase 1 – Assist (Now)**
            - AutoSteer lane keeping, implement guidance
            - Idling reduction & coaching
            - Predictive maintenance, digital work orders
            - Dealer remote visibility

            **Hardware:** Core Hub, GNSS+RTK, AutoSteer HMI, basic sensor tiles
            """
        )
    with col2:
        st.markdown(
            """
            **Phase 2 – Autonomy (Next)**
            - Supervised route execution & implement automation
            - Obstacle detection & control
            - Fleet orchestration + remote supervision
            - Compliance logging

            **Hardware:** Full perception rail, redundant compute, override paddle
            """
        )

    st.session_state.actions["autosteer"] = st.checkbox("Include AutoSteer in ROI", True)
    st.session_state.actions["autonomy"] = st.checkbox("Include Autonomy in ROI", False)

with hardware:
    apply_tab_background("assets/backgrounds/hardware_bg.jpg")
    st.markdown("""
    ### Hardware Vision
    """)
    assets = {
        "Core Module (hero)": "assets/worksmart_core_module.png",
        "Sensor Rail & Tiles": "assets/worksmart_sensor_rail.png",
        "AutoSteer HMI": "assets/worksmart_autosteer_hmi.png",
        "Override Paddle": "assets/worksmart_override_paddle.png",
        "Mounted on Kubota": "assets/worksmart_on_kubota.png",
    }
    colA, colB = st.columns(2)
    for idx, (label, path) in enumerate(assets.items()):
        with (colA if idx < 3 else colB):
            if os.path.exists(path):
                st.image(path, caption=label, use_column_width=True)
            else:
                st.info(f"Missing: {path}")

with anim:
    apply_tab_background("assets/backgrounds/anim_bg.jpg")
    st.markdown("""
    ### Autonomy Animation – Supervised Route
    """)

    field_w, field_h = 100.0, 60.0
    stripe_spacing = 6.0
    speed_mps = st.slider("Vehicle speed (m/s)", 0.5, 5.0, 2.0)
    obstacle_toggle = st.checkbox("Introduce obstacle at mid–field", True)

    stripes = int(field_h // stripe_spacing)
    path = []
    direction = 1
    for i in range(stripes + 1):
        y = i * stripe_spacing
        if direction == 1:
            path += [(0, y), (field_w, y)]
        else:
            path += [(field_w, y), (0, y)]
        direction *= -1
    path = np.array(path)

    obstacle_center = (field_w * 0.55, field_h * 0.5)
    obstacle_r = 5.0
    if obstacle_toggle:
        detour = []
        for (x, y) in path:
            dx, dy = x - obstacle_center[0], y - obstacle_center[1]
            d = (dx * dx + dy * dy) ** 0.5
            if d < obstacle_r:
                off = (dy, -dx)
                norm = max((off[0] ** 2 + off[1] ** 2) ** 0.5, 1e-6)
                scale = (obstacle_r - d) * 0.6
                x, y = x + (off[0] / norm) * scale, y + (off[1] / norm) * scale
            detour.append((x, y))
        path = np.array(detour)

    def interpolate(polyline, step=1.0):
        frames = []
        for i in range(len(polyline) - 1):
            x1, y1 = polyline[i]
            x2, y2 = polyline[i + 1]
            seg = ((x2 - x1) ** 2 + (y2 - y1) ** 2) ** 0.5
            n = max(int(seg / step), 1)
            for t in range(n):
                u = t / n
                frames.append((x1 + u * (x2 - x1), y1 + u * (y2 - y1)))
        frames.append(tuple(polyline[-1]))
        return np.array(frames)

    frames = interpolate(path, step=max(speed_mps * 0.5, 0.2))

    play = st.checkbox("Play", False)
    frame_idx = st.slider("Timeline", 0, len(frames) - 1, 0)

    def draw_frame(idx):
        fig, ax = plt.subplots(figsize=(7, 4))
        ax.add_patch(plt.Rectangle((0, 0), field_w, field_h, fill=False, edgecolor="#777", linewidth=1.5))
        ax.plot(path[:, 0], path[:, 1], color="#1f77b4", linewidth=1.5, alpha=0.7, label="Planned path")
        if obstacle_toggle:
            circ = plt.Circle(obstacle_center, obstacle_r, color="#d62728", alpha=0.2)
            ax.add_patch(circ)
            ax.text(
                obstacle_center[0],
                obstacle_center[1],
                "Obstacle",
                color="#d62728",
                ha="center",
                va="center",
                fontsize=9,
            )
        x, y = frames[idx]
        veh = plt.Circle((x, y), 1.2, color="#2ca02c")
        ax.add_patch(veh)
        if idx < len(frames) - 1:
            x2, y2 = frames[idx + 1]
            ax.arrow(x, y, (x2 - x) * 0.6, (y2 - y) * 0.6, head_width=1.0, head_length=1.5, fc="#2ca02c", ec="#2ca02c")
        ax.set_xlim(-3, field_w + 3)
        ax.set_ylim(-3, field_h + 3)
        ax.set_aspect("equal")
        ax.set_xticks([])
        ax.set_yticks([])
        ax.legend(loc="upper right")
        st.pyplot(fig)

    pct_prog = frame_idx / (len(frames) - 1)
    dist_total = float(np.sum(np.sqrt(np.sum(np.diff(path, axis=0) ** 2, axis=1)))) if len(path) > 1 else 0.0
    est_time = dist_total / max(speed_mps, 0.1)
    st.caption(f"Progress: {pct_prog*100:.0f}% • Distance: {dist_total:.0f} m • ETA @ {speed_mps:.1f} m/s: {est_time/60:.1f} min")

    if play:
        import time as _t
        start = frame_idx
        for i in range(start, len(frames)):
            draw_frame(i)
            _t.sleep(0.05)
            st.experimental_rerun()
    else:
        draw_frame(frame_idx)

with dealer:
    apply_tab_background("assets/backgrounds/dealer_bg.jpg")
    st.markdown("""
    ### Dealer View – Attach Pipeline
    """)
    np.random.seed(0)
    dealers = [f"Dealer {i}" for i in range(1, 11)]
    pipeline = pd.DataFrame(
        {
            "Dealer": dealers,
            "Leads": np.random.randint(20, 120, size=10),
            "Installs Scheduled": np.random.randint(5, 40, size=10),
            "Completed Installs": np.random.randint(0, 35, size=10),
        }
    )
    st.dataframe(pipeline)
    fig, ax = plt.subplots(figsize=(8, 3))
    ax.bar(pipeline["Dealer"], pipeline["Completed Installs"], label="Completed")
    ax.bar(
        pipeline["Dealer"],
        pipeline["Installs Scheduled"],
        bottom=pipeline["Completed Installs"],
        label="Scheduled",
    )
    ax.set_ylabel("Units")
    ax.legend()
    plt.setp(ax.get_xticklabels(), rotation=45, ha="right")
    st.pyplot(fig)

with console:
    apply_tab_background("assets/backgrounds/console_bg.jpg")
    st.markdown("""
    ### Supervision Console
    """)
    rng = np.random.default_rng(42)
    statuses = ["Idle", "Running", "Manual Override", "Error"]
    machines = [f"Unit-{i:03d}" for i in range(1, 13)]

    if st.button("Refresh statuses"):
        st.session_state._fleet = None

    if "_fleet" not in st.session_state or st.session_state._fleet is None:
        st.session_state._fleet = pd.DataFrame(
            {
                "Machine": machines,
                "Status": rng.choice(statuses, size=len(machines), p=[0.2, 0.6, 0.15, 0.05]),
                "Battery%": rng.integers(40, 100, size=len(machines)),
                "Alerts": rng.choice(["", "Obstacle", "GNSS", "Overheat"], size=len(machines), p=[0.7, 0.1, 0.1, 0.1]),
            }
        )
    st.dataframe(st.session_state._fleet)
    st.info("Remote actions: Pause • Resume • Return-to-Base (storyboard only)")

with export:
    apply_tab_background("assets/backgrounds/roi_bg.jpg")
    st.markdown("""
    ### Exports – Executive One–Pager
    """)
    actions = st.session_state.actions
    results = compute_roi(st.session_state.inputs, actions)

    bullets = [
        f"Fleet attached: {results['attached_machines']} of {total_machines(st.session_state.inputs)}",
        f"Gross annual benefit: {currency(results['gross_annual_benefit'])}",
        f"Subscription cost: {currency(results['subscription_cost'])}",
        f"Hardware (Year 0): {currency(results['hardware_cost'])}",
        f"Net annual benefit: {currency(results['net_annual_benefit'])}",
        f"NPV ({st.session_state.inputs['time_horizon_years']} yrs @ {pct(st.session_state.inputs['discount_rate'])}): {currency(results['npv'])}",
    ]
    pm = results["payback_months"]
    bullets.append(f"Payback: {pm:.1f} months" if math.isfinite(pm) else "Payback: > horizon")

    colP, colD = st.columns(2)
    with colP:
        if st.button("Generate PDF"):
            buf = io.BytesIO()
            c = canvas.Canvas(buf, pagesize=LETTER)
            width, height = LETTER
            y = height - 1 * inch

            c.setFont("Helvetica-Bold", 16)
            c.drawString(1 * inch, y, "WorkSmart Hub – Executive Summary")
            y -= 0.3 * inch

            c.setFont("Helvetica", 10)
            c.drawString(1 * inch, y, f"Date: {date.today().isoformat()}")
            y -= 0.3 * inch

            c.setFont("Helvetica", 11)
            for b in bullets:
                c.drawString(1 * inch, y, f"• {b}")
                y -= 0.25 * inch
                if y < 1 * inch:
                    c.showPage()
                    y = height - 1 * inch

            c.showPage()
            c.save()

            st.download_button(
                "Download PDF",
                data=buf.getvalue(),
                file_name="WorkSmart_Executive_OnePager.pdf",
                mime="application/pdf",
            )

    with colD:
        if st.button("Generate Word (.docx)"):
            doc = Document()
            doc.add_heading("WorkSmart Hub – Executive Summary", 0)
            doc.add_paragraph(f"Date: {date.today().isoformat()}")
            for b in bullets:
                doc.add_paragraph(b, style="List Bullet")
            out = io.BytesIO()
            doc.save(out)
            st.download_button(
                "Download DOCX",
                data=out.getvalue(),
                file_name="WorkSmart_Executive_OnePager.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    st.markdown("---")

    st.caption(f"© {date.today().year} Kubota – WorkSmart HUD demo. Synthetic data for illustration only.")
